"""
文件解析服务
支持Word(.docx)和PDF文件的内容提取

注意：
- 用户上传的文件仅在检测期间暂存，检测完成后自动删除
- 不持久化存储用户文稿，保护隐私
"""

import os
import traceback
from config import config


class FileParser:
    """
    文件解析器
    支持格式：.docx (Word), .pdf (PDF), .txt (纯文本)
    """

    # 支持的文件类型
    SUPPORTED_EXTENSIONS = {"docx", "pdf", "txt", "doc"}

    @classmethod
    def parse(cls, filepath: str, original_filename: str) -> dict:
        """
        解析上传的文件，提取文本内容

        参数:
            filepath: 文件的临时存储路径
            original_filename: 原始文件名

        返回:
            dict: {
                "success": bool,
                "text": str,          # 提取的文本内容
                "word_count": int,    # 字数和
                "error": str,         # 错误信息（如有）
                "file_type": str,     # 文件类型
            }
        """
        # 获取文件扩展名
        ext = original_filename.rsplit(".", 1)[-1].lower() if "." in original_filename else ""

        if ext not in cls.SUPPORTED_EXTENSIONS:
            return {
                "success": False,
                "text": "",
                "word_count": 0,
                "error": f"不支持的文件格式：.{ext}，请上传 .docx / .pdf / .txt 文件",
                "file_type": ext,
            }

        try:
            if ext == "docx":
                return cls._parse_docx(filepath, original_filename)
            elif ext == "pdf":
                return cls._parse_pdf(filepath, original_filename)
            elif ext in ("txt", "doc"):
                return cls._parse_txt(filepath, original_filename)
        except Exception as e:
            traceback.print_exc()
            return {
                "success": False,
                "text": "",
                "word_count": 0,
                "error": f"文件解析失败：{str(e)}，请确认文件未损坏",
                "file_type": ext,
            }

    @classmethod
    def _parse_docx(cls, filepath: str, filename: str) -> dict:
        """
        解析Word文档(.docx)
        使用python-docx库提取段落文本
        """
        try:
            from docx import Document
        except ImportError:
            return {
                "success": False, "text": "", "word_count": 0,
                "error": "缺少python-docx依赖，请运行: pip install python-docx",
                "file_type": "docx",
            }

        doc = Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)

        full_text = "\n".join(paragraphs)

        if not full_text.strip():
            return {
                "success": False, "text": "", "word_count": 0,
                "error": "文档内容为空，请检查文件",
                "file_type": "docx",
            }

        from utils.helpers import count_chinese_words
        word_count = count_chinese_words(full_text)

        return {
            "success": True,
            "text": full_text,
            "word_count": word_count,
            "error": "",
            "file_type": "docx",
        }

    @classmethod
    def _parse_pdf(cls, filepath: str, filename: str) -> dict:
        """
        解析PDF文档 - 多引擎 + OCR全覆盖
        策略：PyMuPDF → pdfplumber → PyPDF2 → OCR
        自动选择最佳结果，扫描版图片PDF也能识别
        """
        from utils.helpers import count_chinese_words

        engines_results = []

        # ========== 引擎1：PyMuPDF（最强，支持最多格式）==========
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            total_pages = len(doc)
            pages_text = []
            for page in doc:
                text = page.get_text()
                if text and text.strip():
                    pages_text.append(text.strip())
            doc.close()
            text = "\n".join(pages_text)
            wc = count_chinese_words(text)
            if wc > 5:
                engines_results.append(("PyMuPDF", text, wc))
        except Exception:
            pass

        # ========== 引擎2：pdfplumber（复杂排版友好）==========
        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                pages_text = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        pages_text.append(text.strip())
            text = "\n".join(pages_text)
            wc = count_chinese_words(text)
            if wc > 5:
                # 如果pdfplumber结果更好，替换
                engines_results.append(("pdfplumber", text, wc))
        except Exception:
            pass

        # ========== 引擎3：PyPDF2（最基础）==========
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            # 检查加密
            if reader.is_encrypted:
                return {
                    "success": False, "text": "", "word_count": 0,
                    "error": "PDF已加密，请先解除密码保护",
                    "file_type": "pdf",
                }
            pages_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    pages_text.append(text.strip())
            text = "\n".join(pages_text)
            wc = count_chinese_words(text)
            if wc > 5:
                engines_results.append(("PyPDF2", text, wc))
        except Exception:
            pass

        # ========== 引擎4：OCR（扫描版图片PDF的最后手段）==========
        ocr_text = ""
        for attempt in range(2):
            try:
                from pdf2image import convert_from_path
                import pytesseract
                from PIL import Image

                # 转换PDF为图片
                images = convert_from_path(filepath, dpi=200, first_page=1, last_page=20)
                ocr_pages = []
                for img in images:
                    # 中英文混合OCR
                    page_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                    if page_text and page_text.strip():
                        ocr_pages.append(page_text.strip())
                ocr_text = "\n".join(ocr_pages)
                wc = count_chinese_words(ocr_text)
                if wc > 5:
                    engines_results.append(("OCR识别", ocr_text, wc))
                break
            except Exception:
                if attempt == 0:
                    # 第一次失败，尝试只用英文OCR
                    try:
                        from pdf2image import convert_from_path
                        import pytesseract
                        images = convert_from_path(filepath, dpi=150, first_page=1, last_page=20)
                        ocr_pages = []
                        for img in images:
                            page_text = pytesseract.image_to_string(img, lang='eng')
                            if page_text and page_text.strip():
                                ocr_pages.append(page_text.strip())
                        ocr_text = "\n".join(ocr_pages)
                        wc = count_chinese_words(ocr_text)
                        if wc > 5:
                            engines_results.append(("OCR(英文)", ocr_text, wc))
                    except Exception:
                        pass

        # ========== 选择最佳结果 ==========
        if engines_results:
            # 取字数最多的结果
            best = max(engines_results, key=lambda x: x[2])
            engine_name, full_text, word_count = best

            return {
                "success": True,
                "text": full_text,
                "word_count": word_count,
                "error": "",
                "file_type": "pdf",
            }

        # ========== 全部失败 ==========
        return {
            "success": False,
            "text": ocr_text if ocr_text else "",
            "word_count": count_chinese_words(ocr_text) if ocr_text else 0,
            "error": "所有PDF解析引擎均无法提取文字\n\n"
                     "可能原因：\n"
                     "1. PDF是纯图片扫描版且OCR未安装\n"
                     "2. PDF文件已损坏\n"
                     "3. 建议：用WPS打开 → 另存为.docx → 上传Word文件",
            "file_type": "pdf",
        }

    @classmethod
    def _parse_txt(cls, filepath: str, filename: str) -> dict:
        """
        解析纯文本文件(.txt)
        尝试多种编码方式读取
        """
        # 尝试不同编码读取
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        full_text = ""

        for encoding in encodings:
            try:
                with open(filepath, "r", encoding=encoding) as f:
                    full_text = f.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        if not full_text.strip():
            return {
                "success": False, "text": "", "word_count": 0,
                "error": "文件内容为空或编码不支持",
                "file_type": "txt",
            }

        from utils.helpers import count_chinese_words
        word_count = count_chinese_words(full_text)

        return {
            "success": True,
            "text": full_text,
            "word_count": word_count,
            "error": "",
            "file_type": "txt",
        }
